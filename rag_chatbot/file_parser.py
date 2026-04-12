"""
File text extraction for uploaded attachments.

Supports .txt, .pdf, and .docx. Returns a plain string for injection
into the RAG context. Raises ValueError with a human-readable message
on unsupported types or parse failures.
"""

import io

from fastapi import UploadFile

_MAX_CHARS = 12_000  # ~3 000 tokens — generous but bounded


async def extract_text(file: UploadFile) -> str:
    """Read an uploaded file and return its text content."""
    raw = await file.read()
    filename = (file.filename or "").lower()

    # Determine type by content_type first, fall back to extension
    ct = (file.content_type or "").lower()

    if ct == "text/plain" or filename.endswith(".txt"):
        return _from_txt(raw)

    if ct == "application/pdf" or filename.endswith(".pdf"):
        return _from_pdf(raw)

    if (
        ct == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or filename.endswith(".docx")
    ):
        return _from_docx(raw)

    raise ValueError(
        f"Unsupported file type '{file.content_type}'. "
        "Please upload a .txt, .pdf, or .docx file."
    )


# ──────────────────────────────────────────────────────────────────────────────
# Format handlers
# ──────────────────────────────────────────────────────────────────────────────

def _from_txt(raw: bytes) -> str:
    for encoding in ("utf-8", "latin-1"):
        try:
            return _truncate(raw.decode(encoding).strip())
        except UnicodeDecodeError:
            continue
    raise ValueError("Could not decode the text file. Make sure it is UTF-8 or Latin-1 encoded.")


def _from_pdf(raw: bytes) -> str:
    try:
        import pdfplumber
    except ImportError:
        raise ValueError("PDF support is unavailable (pdfplumber not installed).")

    try:
        pages = []
        with pdfplumber.open(io.BytesIO(raw)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if not text:
                    words = page.extract_words()
                    if words:
                        text = " ".join(w["text"] for w in words)
                if text and text.strip():
                    pages.append(text.strip())

        if pages:
            return _truncate("\n\n".join(pages))

        # No text layer found — fall back to OCR
        return _from_pdf_ocr(raw)

    except ValueError:
        raise
    except Exception as e:
        raise ValueError(f"Failed to read PDF: {e}") from e


def _from_pdf_ocr(raw: bytes) -> str:
    """OCR fallback for scanned / image-based PDFs using pypdfium2 + pytesseract."""
    try:
        import pypdfium2 as pdfium
        import pytesseract
        from PIL import Image  # noqa: F401 — confirms Pillow is available
    except ImportError as e:
        raise ValueError(
            f"OCR support unavailable ({e}). "
            "Install pytesseract and the Tesseract binary to read scanned PDFs."
        )

    # Point pytesseract at the default Windows install location if not on PATH
    import shutil
    if not shutil.which("tesseract"):
        pytesseract.pytesseract.tesseract_cmd = (
            r"C:\Program Files\Tesseract-OCR\tesseract.exe"
        )

    try:
        pdf = pdfium.PdfDocument(raw)
    except Exception as e:
        raise ValueError(f"Failed to open PDF for OCR: {e}") from e

    pages = []
    for i in range(len(pdf)):
        page = pdf[i]
        # scale=2 renders at 144 DPI — sharp enough for reliable OCR
        bitmap = page.render(scale=2)
        pil_image = bitmap.to_pil()
        text = pytesseract.image_to_string(pil_image)
        if text.strip():
            pages.append(text.strip())

    if not pages:
        raise ValueError(
            "OCR found no readable text in the PDF. "
            "The file may be too low-resolution or contain only graphics."
        )

    return _truncate("\n\n".join(pages))


def _from_docx(raw: bytes) -> str:
    try:
        from docx import Document
    except ImportError:
        raise ValueError("DOCX support is unavailable (python-docx not installed).")

    try:
        doc = Document(io.BytesIO(raw))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        if not paragraphs:
            raise ValueError("No text content found in the DOCX file.")
        return _truncate("\n\n".join(paragraphs))
    except ValueError:
        raise
    except Exception as e:
        raise ValueError(f"Failed to read DOCX: {e}") from e


def _truncate(text: str) -> str:
    if len(text) <= _MAX_CHARS:
        return text
    return text[:_MAX_CHARS] + "\n\n... [document truncated — only the first portion was used]"
